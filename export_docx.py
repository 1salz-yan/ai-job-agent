#!/usr/bin/env python3
"""Export als .docx — CV im 2-Spalten-Layout, Anschreiben klassisch."""
import json, os, re, sqlite3
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DEFAULT_EXPORT_DIR = Path.home() / "Desktop" / "Bewerbung" / "Bewerbung"


def _export_dir() -> Path:
    """Export base dir — configurable via settings (export_dir), env EXPORT_DIR,
    falls back to ~/Desktop/Bewerbung/Bewerbung."""
    env = os.getenv("EXPORT_DIR", "").strip()
    if env:
        return Path(env).expanduser()
    try:
        import sqlite3
        dbp = sqlite3.connect(Path(__file__).parent / "job_agent.db")
        row = dbp.execute("SELECT value FROM settings WHERE key='export_dir'").fetchone()
        if row and row[0] and str(row[0]).strip():
            return Path(str(row[0]).strip()).expanduser()
    except Exception:
        pass
    return DEFAULT_EXPORT_DIR
FONT = "Calibri"


def _role_slug(job: dict) -> str:
    """Extract role slug from title (max 2 meaningful words)."""
    title = (job.get("title") or "").strip()
    stopwords = {"bewerbung", "praktikum", "praktikant", "praktikantin", "werkstudent",
                 "student", "junior", "m/w/d", "wm/d", "all", "genders", "interne",
                 "stellenanzeige", "job", "id", "ref",
                 "für", "im", "in", "der", "die", "das", "und", "von", "zum", "zur",
                 "als", "mit", "des", "den", "dem", "ein", "eine", "einer", "eines"}
    # Split on any separator including / and \, strip punctuation & symbols
    import re as _re
    words = [w for w in _re.split(r"[\s/\\|]+", title) if w]
    keep = []
    for w in words:
        w = w.strip("()[]-–—:;.,'\"/\\")
        if len(w) >= 2 and w.lower() not in stopwords:  # skip single letters (m/w/d etc.)
            keep.append(w)
        if len(keep) == 2:
            break
    if not keep:  # fallback: first word ≥2 chars even if it's a stopword
        for w in words:
            w = w.strip("()[]-–—:;.,'\"/\\")
            if len(w) >= 2:
                keep.append(w)
                break
    return "_" + "_".join(keep) if keep else ""


COMPANY_SHORT = {
    "PricewaterhouseCoopers": "PwC",
    "Rhenus Office Systems": "Rhenus",
    "Mercedes-Benz": "Mercedes-Benz",  # keep
}


def _profile_name() -> str:
    """Read display name from local profile DB (falls back to generic)."""
    try:
        import sqlite3
        dbp = sqlite3.connect(Path(__file__).parent / "job_agent.db")
        row = dbp.execute("SELECT name FROM profile WHERE id=1").fetchone()
        return (row[0] or "Name").strip() if row else "Name"
    except Exception:
        return "Name"


def _short_company(name: str) -> str:
    """Abbreviate long company names for filenames (portal limit ~50 chars)."""
    import re as _re
    n = (name or "").strip()
    if len(n) > 20:
        for full, short in COMPANY_SHORT.items():
            if full.lower() in n.lower():
                n = n.replace(full, short)
                break
        # Drop legal suffix + filler tokens if still long
        n = _re.sub(r"\s*(GmbH|AG|SE|KGaA|WPG|Group|Recruiting)\b", "", n, flags=_re.I).strip()
        n = _re.sub(r"\s+", "_", n).strip("_")
        if len(n) > 24:
            n = "_".join(n.split("_")[:2])
    return n


def _folder(job: dict) -> Path:
    """Create folder: YYMMDD_Company_City[_Role] — Role-Teil verhindert
    Überschreiben bei mehreren Bewerbungen am selben Tag bei derselben Firma."""
    now = datetime.now().strftime("%y%m%d")
    company = _short_company(job.get("company") or "Firma").strip().replace(" ", "_")
    loc = (job.get("location") or "").strip()
    city = loc.split(",")[-1].strip() if "," in loc else loc
    city = (city or "Ort").replace(" ", "_")
    path = _export_dir() / f"{now}_{company}_{city}{_role_slug(job)}"
    path.mkdir(parents=True, exist_ok=True)
    return path


def _make_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.15
    for s in doc.sections:
        s.top_margin = Cm(1.5)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(1.8)
        s.right_margin = Cm(1.8)
    return doc


def _label(doc, text: str, size=Pt(9), bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _remove_borders(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                                  '<w:top w:val="nil"/><w:left w:val="nil"/>'
                                  '<w:bottom w:val="nil"/><w:right w:val="nil"/>'
                                  '</w:tcBorders>')
            tcPr.append(tcBorders)


# ============================================================
# CV — 2-Spalten-Layout (Daten kommen aus dem Profil)
# ============================================================
def export_lebenslauf(job: dict, content: str) -> Path:
    """Parse AI output, layout as two-column table CV."""
    import sqlite3
    folder = _folder(job)
    company = _short_company(job.get("company") or "Firma").strip().replace(" ", "_")
    path = folder / f"CV_{_profile_name().replace(' ', '_')}_{company}{_role_slug(job)}.docx"

    # Load profile for left-column data
    db = sqlite3.connect(Path(__file__).parent / "job_agent.db")
    db.row_factory = sqlite3.Row
    pr = db.execute("SELECT * FROM profile WHERE id=1").fetchone()
    profile = {k: pr[k] for k in pr.keys()} if pr else {}
    # Parse AI content into sections
    sections = {}
    current = None
    buf = []
    for line in content.strip().split("\n"):
        s = line.strip()
        if not s or s.startswith("---") or s.startswith("```"):
            continue
        # Clean markdown
        s = s.replace("**", "").replace("## ", "").replace("# ", "")
        upper = s.upper().strip()
        if upper in ("PROFIL", "PROFIL & ZIELE", "SUMMARY", "PROFILE"):
            if current: sections[current] = buf
            current = "profil"; buf = []; continue
        if upper in ("AUSBILDUNG", "AUSBILDUNG & STUDIUM", "EDUCATION", "EDUCATION & TRAINING"):
            if current: sections[current] = buf
            current = "ausbildung"; buf = []; continue
        if upper in ("PRAKTISCHE ERFAHRUNG", "BERUFSERFAHRUNG", "PRAXISERFAHRUNG", "ERFAHRUNG",
                     "EXPERIENCE", "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE", "EXPERIENCE & TRAINING"):
            if current: sections[current] = buf
            current = "erfahrung"; buf = []; continue
        if upper in ("PROJEKTE", "PROJECTS"):
            if current: sections[current] = buf
            current = "projekte"; buf = []; continue
        if upper in ("SPRACHEN & KENNTNISSE", "SPRACHEN UND KENNTNISSE", "SPRACHEN & TOOLS", "KENNTNISSE",
                     "SKILLS", "SKILLS & LANGUAGES", "LANGUAGES & SKILLS", "LANGUAGES", "LANGUAGES & TOOLS"):
            if current: sections[current] = buf
            current = "sprachen"; buf = []; continue
        if current:
            buf.append(s)
    if current:
        sections[current] = buf

    profil_text = "\n".join(sections.get("profil", []))
    if not profil_text:
        profil_text = (profile.get("summary") or "")

    # ---- Build doc ----
    doc = _make_doc()

    # Name (full width at top) — from profile DB (not hardcoded)
    dbp = sqlite3.connect(Path(__file__).parent / "job_agent.db")
    pname = "Name"
    try:
        row = dbp.execute("SELECT name, contact FROM profile WHERE id=1").fetchone()
        if row and row[0]:
            pname = row[0]
    except Exception:
        pass
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(pname)
    run.font.name = FONT
    run.font.size = Pt(18)
    run.bold = True
    name_p.paragraph_format.space_after = Pt(6)

    # Profil — centered below name
    if profil_text:
        prof = doc.add_paragraph()
        prof.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = prof.add_run(profil_text)
        run.font.name = FONT; run.font.size = Pt(9)
        prof.paragraph_format.space_after = Pt(10)

    doc.add_paragraph()

    # ---- 2-column table ----
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    _remove_borders(table)

    # Column widths: left 4.5cm, right 12.5cm
    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(12.5)

    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]

    # ---- LEFT COLUMN ----
    gray = RGBColor(0x66, 0x66, 0x66)

    # Section: Kontakt
    p = left.add_paragraph()
    run = p.add_run("Kontakt")
    run.font.name = FONT; run.font.size = Pt(9); run.bold = True; run.font.color.rgb = gray
    p.paragraph_format.space_after = Pt(3)

    # Contact data — from profile DB (email/address/phone parsed from contact field)
    contact = ""
    try:
        row2 = dbp.execute("SELECT contact FROM profile WHERE id=1").fetchone()
        contact = row2[0] or "" if row2 else ""
    except Exception:
        pass
    parts = [p.strip() for p in contact.split("|")] if contact else []
    email = next((p for p in parts if "@" in p), "")
    phone = next((p for p in parts if p.startswith("+") or p[:1].isdigit() and any(c.isdigit() for c in p)), "")
    addr = next((p for p in parts if not p.startswith("+") and "@" not in p), "")
    for label, val in [("E-Mail", email), ("Adresse", addr), ("Telefon", phone)]:
        p = left.add_paragraph()
        run = p.add_run(f"{label}  ")
        run.font.name = FONT; run.font.size = Pt(8); run.bold = True
        run = p.add_run(val)
        run.font.name = FONT; run.font.size = Pt(8)
        p.paragraph_format.space_after = Pt(1)

    # Section: Sprachen
    left.add_paragraph()
    p = left.add_paragraph()
    run = p.add_run("Sprachen")
    run.font.name = FONT; run.font.size = Pt(9); run.bold = True; run.font.color.rgb = gray
    p.paragraph_format.space_after = Pt(3)
    langs = profile.get("languages", "").split("·") if profile.get("languages") else []
    for l in langs:
        l = l.strip()
        if not l:
            continue
        if "(" in l:
            lang, level = l.split("(", 1)
            lang = lang.strip(); level = level.strip(") ")
        else:
            lang = l; level = ""
        p = left.add_paragraph()
        run = p.add_run(lang)
        run.font.name = FONT; run.font.size = Pt(8); run.bold = True
        if level:
            run = p.add_run(f"\n{level}")
            run.font.name = FONT; run.font.size = Pt(8); run.font.color.rgb = gray
        p.paragraph_format.space_after = Pt(2)

    # Section: Tools & Kenntnisse
    left.add_paragraph()
    p = left.add_paragraph()
    run = p.add_run("Tools & Kenntnisse")
    run.font.name = FONT; run.font.size = Pt(9); run.bold = True; run.font.color.rgb = gray
    p.paragraph_format.space_after = Pt(3)
    skills = profile.get("skills", "").split("·") if profile.get("skills") else []
    for s in skills[:12]:
        s = s.strip()
        if not s:
            continue
        p = left.add_paragraph()
        run = p.add_run(s)
        run.font.name = FONT; run.font.size = Pt(8)
        p.paragraph_format.space_after = Pt(1)

    # ---- RIGHT COLUMN ----
    # Profil is already rendered above table, not repeated here
    pass

    # Section helper
    def right_section(title, lines):
        # Section header with I prefix
        p = right.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(f"I  {title}")
        run.font.name = FONT; run.font.size = Pt(11); run.bold = True; run.font.color.rgb = gray
        p.paragraph_format.space_after = Pt(4)

        # Parse entries
        entries = []
        current_entry = []
        for line in lines:
            s = line.strip()
            if not s:
                continue
            # New entry: starts without "-" and contains "|"
            if not s.startswith("-") and "|" in s:
                if current_entry:
                    entries.append(current_entry)
                current_entry = [s]
            else:
                current_entry.append(s)
        if current_entry:
            entries.append(current_entry)

        for entry in entries:
            header = entry[0].lstrip("- ")
            p = right.add_paragraph()
            run = p.add_run(header)
            run.font.name = FONT; run.font.size = Pt(9.5); run.bold = True
            p.paragraph_format.space_after = Pt(1)

            for bullet in entry[1:]:
                b = bullet.lstrip("- ").strip()
                if not b:
                    continue
                p = right.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(f"• {b}")
                run.font.name = FONT; run.font.size = Pt(9)
                p.paragraph_format.space_after = Pt(1)

    # Render right-column sections
    # (2026-08: removed the old keep() filter that dropped projects containing
    # drone/blockchain/crypto/solana when the JD didn't mention them — it
    # overrode the AI's job-specific project selection and silently deleted
    # projects like Machine Economy (DePIN/Solana) from the export. The AI
    # already picks projects per JD via 'Passt zu' tags; export must be WYSIWYG.)
    p_raw = sections.get("projekte", [])
    filtered_proj = []
    buf = []
    for s in p_raw:
        if s and not s.startswith("-") and "|" in s:
            if buf: filtered_proj.extend(buf)
            buf = [s]
        elif s:
            buf.append(s)
    if buf: filtered_proj.extend(buf)

    right_section("Berufserfahrung", sections.get("erfahrung", []))
    right_section("Ausbildung", sections.get("ausbildung", []))
    right_section("Projekte", filtered_proj)

    doc.save(str(path))
    return path


# ============================================================
# ANSCHREIBEN — klassisch, wie Vorlage
# ============================================================
def export_anschreiben(job: dict, content: str) -> Path:
    folder = _folder(job)
    company = _short_company(job.get("company") or "Firma").strip().replace(" ", "_")
    path = folder / f"Anschreiben_{_profile_name().replace(' ', '_')}_{company}{_role_slug(job)}.docx"

    loc = (job.get("location") or "").strip()
    city = loc.split(",")[-1].strip() if "," in loc else loc

    doc = _make_doc()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2)

    def line(text, size=Pt(11), bold=False, space=0):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = FONT; run.font.size = size
        if bold: run.bold = True
        p.paragraph_format.space_after = Pt(space)

    # Sender — address/contact from profile DB
    line(_profile_name(), size=Pt(13), bold=True)
    try:
        import sqlite3 as _sq
        _c = ""
        _r = _sq.connect(Path(__file__).parent / "job_agent.db").execute("SELECT contact FROM profile WHERE id=1").fetchone()
        _c = _r[0] or "" if _r else ""
        _parts = [p.strip() for p in _c.split("|")] if _c else []
        _addr = next((p for p in _parts if "@" not in p and not p.startswith("+")), "")
        _email = next((p for p in _parts if "@" in p), "")
        _phone = next((p for p in _parts if p.startswith("+")), "")
        line(_addr or "Adresse")
        line((_email + "  |  " + _phone).strip(" |"))
    except Exception:
        line("Adresse")
        line("E-Mail  |  Telefon")
    doc.add_paragraph()

    # Recipient
    line(company)
    line("z.Hd. Personalabteilung")
    line(city)
    doc.add_paragraph()

    # Date — kein führendes Null (DIN/karrierebibel: "4.5.2026")
    line(f"Berlin, {datetime.now().strftime('%-d.%-m.%Y')}")
    doc.add_paragraph()

    # Subject
    lines = [l.strip() for l in content.strip().split("\n") if l.strip()]
    lines = [l.replace("**", "") for l in lines]  # strip markdown
    betreff = ""
    body_start = 0
    # Subject — keine Markdown-Sterne, kein Wort "Betreff" (karrierebibel)
    # AI liefert erste Zeile als "Bewerbung: ..." — das ist der Betreff
    if lines and re.match(r"^Bewerbung\b", lines[0]):
        betreff = lines[0]
        body_start = 1
    else:
        for i, l in enumerate(lines):
            if l.lower().startswith("betreff"):
                betreff = re.sub(r"^[Bb]etreff\s*[:–-]?\s*", "", l).strip()
                body_start = i + 1; break
    if not betreff and job.get("title"):
        betreff = f"Bewerbung: {job['title']}"
    if betreff:
        line(betreff, bold=True, space=4)
        doc.add_paragraph()

    # Body — strip AI-generated closing to avoid duplication
    body = "\n".join(lines[body_start:]).strip() if body_start else content.strip()
    # Remove everything from first "Mit freundlichen Grüßen" onward
    for cutoff in ["Mit freundlichen Grüßen", "Mit freundlichem Gruß"]:
        idx = body.find(cutoff)
        if idx >= 0:
            body = body[:idx].strip()
            # Also strip trailing Anlagen/Name lines
            for trailer in ["\n\nAnlagen:", "\nAnlagen:", f"\n\n{_profile_name()}", f"\n{_profile_name()}"]:
                tidx = body.rfind(trailer)
                if tidx > len(body) - 80:  # only near the end
                    body = body[:tidx].strip()
            break
    for para in body.split("\n\n"):
        p = para.strip()
        if p:
            pp = doc.add_paragraph()
            pp.paragraph_format.space_after = Pt(10)
            run = pp.add_run(p)
            run.font.name = FONT; run.font.size = Pt(11)

    doc.add_paragraph()
    line("Mit freundlichen Grüßen")  # kein Komma nach Grußformel (DIN/karrierebibel)
    doc.add_paragraph()
    line(_profile_name())
    doc.add_paragraph()
    line("Anlagen: Lebenslauf, Immatrikulationsbescheinigung, Notenspiegel, relevante Zeugnisse")

    doc.save(str(path))
    return path
